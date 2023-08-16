from docx import Document
from tkinter import font
from docx.shared import Pt
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkinter import filedialog, messagebox
import tkinter as tk
from tkinter import ttk
from ttkthemes import ThemedTk
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os



def extract_text_from_docx(file_path, keywords):
    doc = Document(file_path)
    full_text = [paragraph.text for paragraph in doc.paragraphs]
    text = '\n'.join(full_text)
    found_keywords = []

    for keyword in keywords:
        if keyword in text:
            found_keywords.append(keyword)

    return found_keywords, text


def wrap_text(text, max_line_length):
    words = text.split()
    lines = []
    space_padding = "                    "  # This is a string with spaces
    current_line = space_padding  # Start with the padding

    for word in words:
        # Check if adding the word to the current line would exceed the max line length
        if len(current_line + word + " ") <= max_line_length:
            current_line += word + " "
        else:
            lines.append(current_line.strip())  # Append the current line to the list of lines
            current_line = space_padding + word + " "  # Start a new line with the padding, followed by the word

    if current_line.strip():  # If there's any content left in current_line
        lines.append(current_line.strip())

    return '\n'.join(lines)


def select_directory_and_scan(text_widget, status_label):
    print("Function called!")

    keywords = keywords_entry.get().split(',')
    directory = filedialog.askdirectory()

    print(f"Directory selected: {directory}")

    keyword_counts = {keyword: 0 for keyword in keywords}
    files_with_keywords = []

    if directory:
        text_widget.delete(1.0, tk.END)

        for filename in os.listdir(directory):
            print(f"Processing file: {filename}")

            if filename.endswith(".docx"):
                full_path = os.path.join(directory, filename)
                try:
                    found_keywords, extracted_text = extract_text_from_docx(full_path, keywords)

                    if found_keywords:
                        for keyword in found_keywords:
                            keyword_counts[keyword] += 1

                        files_with_keywords.append((filename, found_keywords))

                        wrapped_text = wrap_text(extracted_text, max_line_length=130)  # Adjust line length

                        text_widget.insert(tk.END, f"\n{'-' * 100}\n", ("normal",))
                        text_widget.insert(tk.END, f"File: {filename}\n", ("filename",))
                        text_widget.insert(tk.END, f"Keywords found: {', '.join(found_keywords)}\n", ("normal",))
                        text_widget.insert(tk.END, f"{'-' * 100}\n", ("normal",))
                        text_widget.insert(tk.END, f"{wrapped_text}\n", ("normal",))
                        text_widget.insert(tk.END, f"{'-' * 100}\n\n", ("normal",))

                except Exception as e:
                    print(f"Error processing file {filename}. Error: {e}")

        text_widget.tag_configure("normal", justify="center", font=("Arial", 12))
        text_widget.tag_configure("filename", justify="center", font=("Arial", 16, "bold"))
        text_widget.tag_configure("summary", justify="center", font=("Arial", 20, "bold"), foreground="#008080")

        summary = "\n\nSummary:\n"
        for keyword, count in keyword_counts.items():
            summary += f"We found the keyword '{keyword}' in {count} files:\n"

        text_widget.insert("1.0", summary, ("summary",))

        # Add buttons for copying and saving as Word file
        copy_button = tk.Button(main_frame, text="Copy Output", font=button_font_style,
                                bg="#5E81AC", fg=font_color, activebackground="#88C0D0", relief=tk.FLAT,
                                padx=20, pady=10, command=lambda: copy_output(text_widget))
        copy_button.pack(pady=10)

        save_button = tk.Button(main_frame, text="Save as Word", font=button_font_style,
                                bg="#5E81AC", fg=font_color, activebackground="#88C0D0", relief=tk.FLAT,
                                padx=20, pady=10, command=lambda: save_as_word(text_widget))
        save_button.pack(pady=10)



def copy_output(text_widget):
    output_text = text_widget.get("1.0", tk.END)
    root.clipboard_clear()
    root.clipboard_append(output_text)
    root.update()


def save_as_word(text_widget):
    output_text = text_widget.get("1.0", tk.END)
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])

    if file_path:
        doc = Document()

        # Add extracted text with formatting
        paragraphs = output_text.split('\n')
        for paragraph in paragraphs:
            if paragraph.startswith("-"):
                p = doc.add_paragraph(paragraph)
            elif paragraph.startswith("File:"):
                p = doc.add_paragraph(paragraph)
            elif paragraph.startswith("Keywords found:"):
                p = doc.add_paragraph(paragraph)
            else:
                p = doc.add_paragraph(paragraph)

            # Set paragraph alignment to right-to-left
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Save the document
        doc.save(file_path)


# Set up the main window
root = ThemedTk(theme="equilux")
root.title("Keyword-Based Files Scanner")
root.geometry("1600x1000")

# Define styles and colors
bg_color = "#2E3440"  # Dark background
font_color = "#ECEFF4"  # Light font color
font_style = ("Ariel", 14)
header_font_style = ("Ariel", 20, "bold")
button_font_style = ("Ariel", 14)

# Styling frames and labels for a consistent look
style = ttk.Style()
style.configure('TFrame', background=bg_color)
style.configure('TLabel', background=bg_color, foreground=font_color)

# Main content frame
main_frame = ttk.Frame(root, padding="10")
main_frame.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)  # Reduced padding for top balance

# Display header
header_label = ttk.Label(main_frame, text="Keyword-Based Files Scanner", font=header_font_style)
header_label.pack(pady=10)  # Reduced padding for balance

# User input for keywords
keywords_entry = ttk.Entry(main_frame, font=font_style, width=100)
keywords_entry.insert(0, "Enter keywords separated by commas")
keywords_entry.pack(pady=20)

# Button to initiate file scanning
button = tk.Button(main_frame, text="Select a Folder", font=button_font_style,
                   bg="#5E81AC", fg=font_color, activebackground="#88C0D0", relief=tk.FLAT, padx=20, pady=10,
                   command=lambda: select_directory_and_scan(result_text, status_label))
button.pack(pady=30)

# Create the result_text Text widget with scrollbars
result_text = tk.Text(main_frame, wrap=tk.WORD, height=20, width=110, font=font_style, bg="#3B4252", fg=font_color,
                      spacing2=9)
result_text.pack(padx=10, pady=10, side=tk.LEFT, fill=tk.BOTH, expand=True)

# Scrollbar for result area
scroll = tk.Scrollbar(main_frame, command=result_text.yview, bg="#4C566A")
result_text.configure(yscrollcommand=scroll.set, relief=tk.FLAT)
scroll.pack(side=tk.RIGHT, fill=tk.Y)

# Label to display the status of scanning
status_label = ttk.Label(main_frame, text="", font=font_style)
status_label.pack(pady=20)

# Credit label
credit_label = ttk.Label(root, text="Developed by Maya Rom", font=("Ubuntu", 10, "italic"))
credit_label.pack(side=tk.BOTTOM, pady=5)

root.mainloop()
